import re
from datetime import datetime
from io import BytesIO
from urllib.parse import quote_plus

import streamlit as st
import feedparser

from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import docx.opc.constants as constants

# -----------------------------
# Page config (must be called once, at the top)
# -----------------------------
st.set_page_config(
    page_title="Jerry Moran News Search",
    page_icon="📰",
    layout="wide",
)

# -----------------------------
# Config & constants
# -----------------------------
DEFAULT_TERMS = [
    term + " -Allan -Anna -Terry"
    for term in [
        "Jerry Moran", "Senator Jerry Moran", "Senator Moran",
        "Sen. Moran", "Sen. Jerry Moran", "Sens. Moran", "Sens. Jerry Moran"
    ]
]

KANSAS_OUTLETS = [
    'Kansas Reflector', 'The Topeka Capital-Journal', 'The Wichita Eagle',
    'KCLY Radio', 'KSN-TV', 'KWCH', 'Kansas City Star',
    'Lawrence Journal-World', 'The Garden City Telegram', 'KSNT 27 News',
    'The Hutchinson News', 'Salina Journal', 'Hays Daily News',
    'Hays Post', 'Emporia Gazette', 'JC Post', 'WIBW', 'KRSL', 'Dodge City Daily Globe', 'DerbyInformer.com', 'KCTV'
]

EXCLUDE_SOURCES_CONTAINS = ['.gov', 'Quiver Quantitative', 'MSN', 'Twin States News']

# -----------------------------
# Light CSS to center content and tighten spacing
# -----------------------------
st.markdown(
    """
    <style>
      .block-container { padding-top: 1.2rem; max-width: 1100px; }
      ol li { margin-bottom: .35rem; }
      a { text-decoration: underline; }
      .center-text { text-align:center; }
    </style>
    """,
    unsafe_allow_html=True,
)

# -----------------------------
# Google News RSS helpers (no pygooglenews needed)
# -----------------------------
def google_news_rss(term: str, when: str = "1d", lang="en-US", country="US") -> str:
    """Build a Google News RSS query. Supports query operator `when:` (e.g., 1d, 7d)."""
    query = f"{term} when:{when}"
    q = quote_plus(query)
    return f"https://news.google.com/rss/search?q={q}&hl={lang}&gl={country}&ceid={country}:en"


def fetch_entries(search_terms, when="1d"):
    """Fetch & de-duplicate entries across terms using Google News RSS + feedparser."""
    all_entries, seen_links = [], set()
    for term in search_terms:
        url = google_news_rss(term, when=when)
        feed = feedparser.parse(url)
        for e in feed.entries:
            link = e.get("link")
            if not link or link in seen_links:
                continue
            seen_links.add(link)

            # Source / outlet
            media = None
            src = getattr(e, "source", None)
            if src and isinstance(src, dict):
                media = src.get("title")
            if not media:
                media = e.get("author") or "Unknown"

            title = (e.get("title") or "").strip()
            all_entries.append({
                "title": title,
                "link": link,
                "source": {"title": media},
            })
    return all_entries

# -----------------------------
# Helpers for cleanup, dedupe, and DOCX
# -----------------------------
def clean_text(text: str) -> str:
    return re.sub(r'[^\w\s\-\.,:;!?()\'\"]+', '', text or '').strip()


def normalize_title_for_duplicate_detection(title: str) -> str:
    normalized = (title or '').lower().strip()
    normalized = re.sub(r'^(breaking:?\s*|update:?\s*|exclusive:?\s*)', '', normalized)
    normalized = re.sub(r'\s+', ' ', normalized)
    return normalized


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    color = OxmlElement('w:color'); color.set(qn('w:val'), '0000FF'); rPr.append(color)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)

    new_run.append(rPr)
    text_elem = OxmlElement('w:t'); text_elem.text = text; new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def process_entries_with_duplicates(all_entries, kansas_outlets=None):
    if kansas_outlets is None:
        kansas_outlets = KANSAS_OUTLETS

    title_groups = {}
    for entry in all_entries:
        media = (entry['source']['title'] or "Unknown")
        if any(x in media for x in EXCLUDE_SOURCES_CONTAINS):
            continue

        raw_title = re.sub(rf" - {re.escape(media)}$", "", entry['title'] or "")
        title = clean_text(raw_title)
        normalized = normalize_title_for_duplicate_detection(title)

        title_groups.setdefault(normalized, []).append({
            'title': title,
            'media': clean_text(media),
            'link': entry['link'],
            'entry': entry,
        })

    processed = []
    for group in title_groups.values():
        if not group:
            continue
def process_entries_with_duplicates(all_entries, kansas_outlets=None):
    if kansas_outlets is None:
        kansas_outlets = KANSAS_OUTLETS

    def is_kansas_outlet(media, kansas_outlets):
        return any(k.strip() and k.strip().lower() == media.lower() for k in kansas_outlets)

    title_groups = {}
    for entry in all_entries:
        media = (entry['source']['title'] or "Unknown")
        if any(x in media for x in EXCLUDE_SOURCES_CONTAINS):
            continue

        raw_title = re.sub(rf" - {re.escape(media)}$", "", entry['title'] or "")
        title = clean_text(raw_title)
        normalized = normalize_title_for_duplicate_detection(title)

        title_groups.setdefault(normalized, []).append({
            'title': title,
            'media': clean_text(media),
            'link': entry['link'],
            'entry': entry,
        })

    processed = []
    for group in title_groups.values():
        if not group:
            continue
        primary, duplicates = group[0], group[1:]
        media_string = primary['media']
        if is_kansas_outlet(media_string, kansas_outlets):
            media_string = f"*{media_string}"
        if duplicates:
            def format_outlet(media):
                return f"*{media}" if is_kansas_outlet(media, kansas_outlets) else media

            dup_outlets = [format_outlet(d['media']) for d in duplicates]
            if len(dup_outlets) == 1:
                media_string += f" (also ran in {dup_outlets[0]})"
            else:
                media_string += f" (also ran in {', '.join(dup_outlets[:-1])} and {dup_outlets[-1]})"

        is_kansas = is_kansas_outlet(primary['media'], kansas_outlets)

        processed.append({
            'title': primary['title'],
            'media_string': media_string,
            'link': primary['link'],
            'is_kansas': is_kansas,
        })
    return processed


def build_docx_bytes(processed_entries):
    now = datetime.now()
    filename = f"Sen Moran in the News {now.month}-{now.day}.docx"
    doc = Document()
    doc.add_paragraph(f"Generated on: {now.strftime('%B %d, %Y at %I:%M %p')}")
    doc.add_paragraph(f"Total articles found: {len(processed_entries)}")
    doc.add_paragraph("* indicates Kansas news outlet")
    doc.add_paragraph()
    doc.add_heading('News Articles', level=1)

    for entry in processed_entries:
        # No need for extra star prefix, it's handled in media_string now
        p = doc.add_paragraph()
        p.add_run(f"{entry['media_string']}: ")
        add_hyperlink(p, entry['link'], entry['title'])
        url_run = p.add_run(f" [{entry['link']}]")
        url_run.font.italic = True
        url_run.font.size = Pt(8)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return filename, bio

# -----------------------------
# HEADER (Centered, landscape image)
# -----------------------------
# Try local asset first; fallback to placeholder
try:
    st.image("assets/jerry-moran.jpg", use_column_width=True)
except Exception:
    st.image("https://via.placeholder.com/1200x400?text=Jerry+Moran", use_column_width=True)

st.markdown("<h1 class='center-text'>Jerry Moran — News Tracker</h1>", unsafe_allow_html=True)
st.markdown(
    "<p class='center-text'>Live Google News RSS search with smart deduping, Kansas-outlet highlighting, and one-click DOCX export.</p>",
    unsafe_allow_html=True,
)
st.markdown("<hr>", unsafe_allow_html=True)

# -----------------------------
# SEARCH CONTROLS (Centered)
# -----------------------------
col1, col2, col3 = st.columns([1, 2, 1])
with col2:
    terms_text = st.text_area(
        "Search terms (comma-separated)",
        value=", ".join([
            "Jerry Moran", "Senator Jerry Moran", "Senator Moran",
            "Sen. Moran", "Sen. Jerry Moran", "Sens. Moran", "Sens. Jerry Moran"
        ]),
        height=110,
    )
    exclude_text = st.text_area(
        "Exclude these words (comma-separated)",
        value="Allan, Anna, Terry",
        height=60,
    )
    kansas_media_text = st.text_area(
        "Kansas media outlets (comma-separated)",
        value=", ".join(KANSAS_OUTLETS),
        height=80,
    )
    when_choice = st.selectbox(
        "Time window",
        options=["1d", "3d", "7d", "14d", "30d"],
        index=0,
        help="Google News query operator `when:`",
    )
    run_search = st.button("🔎 Run Search")

diff --git a/app.py b/app.py
index e4de0512e76e3c6e20d4a78f3919842d2d43a17e..3dd09a8f3d7c6f26abca0e14db90346a36c54f9e 100644
--- a/app.py
+++ b/app.py
@@ -280,43 +280,42 @@ with col2:
 
 # -----------------------------
 # RESULTS (Centered column)
 # -----------------------------
 if run_search:
     search_terms = [t.strip() for t in terms_text.split(",") if t.strip()]
     exclude_terms = [e.strip() for e in exclude_text.split(",") if e.strip()]
     kansas_media = [k.strip() for k in kansas_media_text.split(",") if k.strip()]
 
     # Append negative keywords to each search term
     if exclude_terms:
         negatives = " ".join([f"-{word}" for word in exclude_terms])
         search_terms = [term + " " + negatives for term in search_terms]
 
     with st.spinner("Searching Google News…"):
         all_entries = fetch_entries(search_terms, when=when_choice)
 
     # Pass the editable Kansas list directly
     processed_entries = process_entries_with_duplicates(all_entries, kansas_outlets=kansas_media)
 
     st.markdown(
         f"<p class='center-text'><strong>Found {len(all_entries)} items before dedupe — After dedupe: {len(processed_entries)}</strong></p>",
         unsafe_allow_html=True,
     )
 
-# ... inside your Streamlit results display block ...
-c1, c2, c3 = st.columns([0.5, 3, 0.5])
-with c2:
-    md_lines = []
-    for i, entry in enumerate(processed_entries, 1):
-        md_lines.append(f"{i}. {entry['media_string']}: [{entry['title']}]({entry['link']})")
-    st.markdown("\n".join(md_lines))
-
-    filename, bio = build_docx_bytes(processed_entries)
-    st.download_button(
-        "⬇️ Download Word Document",
-        data=bio,
-        file_name=filename,
-        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
-    )
+    # ... inside your Streamlit results display block ...
+    c1, c2, c3 = st.columns([0.5, 3, 0.5])
+    with c2:
+        md_lines = []
+        for i, entry in enumerate(processed_entries, 1):
+            md_lines.append(f"{i}. {entry['media_string']}: [{entry['title']}]({entry['link']})")
+        st.markdown("\n".join(md_lines))
+
+        filename, bio = build_docx_bytes(processed_entries)
+        st.download_button(
+            "⬇️ Download Word Document",
+            data=bio,
+            file_name=filename,
+            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
         )
 else:
     st.info("Enter search terms above and click **Run Search**.")

